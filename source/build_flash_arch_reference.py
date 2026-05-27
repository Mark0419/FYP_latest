from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Users\wnmin\OneDrive\Attachments\FYP_source\Flash_Controller_Architecture_Reference.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths):
    set_table_width(table, widths)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(31, 77, 120)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 58, 95)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Current specification"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    style_table(table, [1.75, 4.75])
    doc.add_paragraph()
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, widths)
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.font.size = Pt(11)

for name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)
run = title.add_run("Flash Controller Architecture Reference")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
r = subtitle.add_run("Legacy SPI SRAM shadowing vs pure QSPI XIP vs dual-mode QSPI")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(85, 85, 85)

add_body(
    doc,
    "This reference captures the current hardware-level test intent for the three flash-controller architectures. "
    "All designs are driven by the Cortex-M3 at SoC level, use 32-bit addressing, and are measured with a common 4-word flash payload."
)

add_heading(doc, "Common Test Baseline", 1)
add_kv_table(doc, [
    ("CPU", "ARM Cortex-M3 integration model"),
    ("Clock", "50 MHz HCLK, 20 ns period"),
    ("Reset", "Held for 20 HCLK cycles, then released on the aligned reset sequence"),
    ("Addressing", "32-bit AHB addresses; QSPI flash accesses use 32-bit address fields"),
    ("AHB data width", "32-bit CPU/system data"),
    ("Common payload size", "4 x 32-bit words = 16 bytes of data"),
    ("Completion marker", "A final write of 0xAABBCCDD to 0x20008000 tells the testbench that the benchmark completed"),
])

add_body(doc, "Common flash payload:")
add_code_block(doc, [
    "0x40000000 = 49024801",
    "0x40000004 = e7fe6001",
    "0x40000008 = 20008000",
    "0x4000000c = aabbccdd",
])

add_matrix_table(
    doc,
    ["Architecture", "Fundamental idea", "Current benchmark behavior", "What timing means"],
    [
        (
            "Legacy SPI + SRAM shadowing",
            "Copy flash code/data into SRAM first, then execute from SRAM.",
            "Bootloader reads 4 words through SPI, writes them to SRAM, jumps to 0x20000001, and the SRAM payload writes the magic marker.",
            "Initialization is shadow-copy time. Operation is SRAM payload execution time after shadowing is ready.",
        ),
        (
            "Pure QSPI XIP",
            "No shadowing; CPU reads directly from memory-mapped QSPI flash.",
            "Firmware performs four direct reads at 0x40000000..0x4000000c. It writes only the final marker to SRAM for testbench completion.",
            "Initialization is first XIP read latency. Operation is remaining direct-read time.",
        ),
        (
            "Dual-mode QSPI",
            "QSPI controller supports XIP reads plus program/erase capability.",
            "For PPA read comparison, only the direct-read/XIP path is exercised. Write/erase path is intentionally not part of this benchmark.",
            "Same timing definition as pure XIP, while area/power include the richer dual-mode hardware.",
        ),
    ],
    [1.45, 1.65, 2.1, 1.3],
)

add_heading(doc, "Architecture 1: Legacy SPI With SRAM Shadowing", 1)
add_body(
    doc,
    "The legacy architecture treats external flash as a boot source rather than the long-term execution target. "
    "The Cortex-M3 first runs a ROM bootloader. The bootloader reads words from flash through a memory-mapped SPI controller, stores them into SRAM, and then jumps into SRAM."
)

add_heading(doc, "Key Specification", 2)
add_kv_table(doc, [
    ("Controller bus", "AHB-Lite slave mapped in the 0x40000000 region"),
    ("Physical interface", "Single-bit SPI: spi_cs_n, spi_clk, spi_mosi, spi_miso"),
    ("Read command", "0x03"),
    ("Address phase", "32-bit flash address"),
    ("Data phase", "32-bit serial receive"),
    ("Destination", "SRAM at 0x20000000"),
    ("Execution jump", "0x20000001, because Cortex-M3 requires bit 0 set for Thumb state"),
])

add_heading(doc, "SPI Register Map", 2)
add_matrix_table(
    doc,
    ["Address", "Register", "Role"],
    [
        ("0x40000000", "SPI_CTRL_REG", "bit[0] start, bit[1] busy"),
        ("0x40000004", "SPI_CMD_REG", "Holds command byte, currently 0x03 for read"),
        ("0x40000008", "SPI_ADDR_REG", "Holds 32-bit flash byte address"),
        ("0x4000000c", "SPI_DATA_REG", "Returns received 32-bit word"),
    ],
    [1.5, 1.5, 3.5],
)

add_heading(doc, "Legacy Test Flow", 2)
add_code_block(doc, [
    "ROM bootloader starts after reset",
    "for i = 0..3:",
    "    write SPI_CMD_REG  = 0x03",
    "    write SPI_ADDR_REG = flash_base + i*4",
    "    write SPI_CTRL_REG = 1",
    "    wait for busy to assert, then deassert",
    "    read SPI_DATA_REG",
    "    write SRAM[0x20000000 + i*4]",
    "jump to 0x20000001",
    "SRAM payload writes 0xaabbccdd to 0x20008000",
])

add_heading(doc, "Legacy Timing Metrics", 2)
add_kv_table(doc, [
    ("Init time", "Reset release to the last SRAM shadow write at 0x2000000c"),
    ("Operation time", "Last SRAM shadow write to the payload's magic write at 0x20008000"),
    ("Total time", "Reset release to the magic write"),
])

add_heading(doc, "Architecture 2: Pure QSPI XIP", 1)
add_body(
    doc,
    "Pure QSPI XIP removes the SRAM shadow step. The CPU accesses a memory-mapped flash window directly, and the AHB wrapper stalls the CPU until the QSPI read completes. "
    "The current benchmark does direct reads only; it does not copy payload words into SRAM."
)

add_heading(doc, "Key Specification", 2)
add_kv_table(doc, [
    ("Controller bus", "AHB-Lite slave mapped in the 0x40000000 flash window"),
    ("Physical interface", "Quad SPI IO[3:0] with qspi_cs_n and qspi_clk"),
    ("Read command", "0xEC on the first read"),
    ("Address phase", "32-bit address transmitted as quad nibbles"),
    ("Mode bits", "0xA5 to enter/maintain continuous XIP mode"),
    ("Dummy cycles", "4 QSPI clocks before data"),
    ("Data phase", "32-bit data received over four QSPI IO lines"),
    ("Write/erase", "Not supported in pure XIP architecture"),
])

add_heading(doc, "Pure XIP FSM Flow", 2)
add_code_block(doc, [
    "First read when xip_active == 0:",
    "    IDLE -> CMD(0xEC) -> ADDR(32-bit) -> MODE(0xA5) -> DUMMY -> DATA -> DONE",
    "Later reads when xip_active == 1:",
    "    IDLE -> ADDR(32-bit) -> MODE(0xA5) -> DUMMY -> DATA -> DONE",
])

add_heading(doc, "Pure XIP Test Flow", 2)
add_code_block(doc, [
    "ROM firmware starts after reset",
    "read 0x40000000",
    "read 0x40000004",
    "read 0x40000008",
    "read 0x4000000c",
    "if final word is 0xaabbccdd, write 0xaabbccdd to 0x20008000",
])

add_heading(doc, "Pure XIP Timing Metrics", 2)
add_kv_table(doc, [
    ("Init time", "Reset release to the first direct XIP read completion"),
    ("Operation time", "First XIP read completion to fourth XIP read completion"),
    ("Total read time", "Reset release to fourth XIP read completion"),
    ("Marker latency", "Fourth read completion to final marker write; useful to exclude from flash-read comparison"),
])

add_heading(doc, "Architecture 3: Dual-Mode QSPI", 1)
add_body(
    doc,
    "Dual-mode QSPI keeps the XIP/direct-read path but adds flash program and erase capability. "
    "For the current PPA read comparison, the benchmark intentionally uses only the direct-read path so it can be compared fairly with pure XIP."
)

add_heading(doc, "Key Specification", 2)
add_kv_table(doc, [
    ("Controller bus", "AHB-Lite slave mapped in the 0x40000000 flash window"),
    ("Physical interface", "Quad SPI IO[3:0] with qspi_cs_n and qspi_clk"),
    ("Read command", "0xEC, same XIP read command as pure XIP"),
    ("Write enable", "0x06"),
    ("Page program", "0x34"),
    ("Sector erase", "0x21"),
    ("Erase selector", "HWDATA == 0xffffffff selects erase in the current RTL"),
    ("XIP behavior", "Reads can skip the command phase after xip_active is locked"),
])

add_heading(doc, "Dual-Mode FSM Flow", 2)
add_code_block(doc, [
    "Read path:",
    "    IDLE -> CMD(0xEC) -> ADDR -> MODE -> DUMMY -> DATA -> DONE",
    "    or, if xip_active == 1:",
    "    IDLE -> ADDR -> MODE -> DUMMY -> DATA -> DONE",
    "Write/program path:",
    "    IDLE -> WREN(0x06) -> CMD(0x34) -> ADDR -> DATA/PROG -> DONE",
    "Erase path:",
    "    IDLE -> WREN(0x06) -> CMD(0x21) -> ADDR -> PROG/DONE",
])

add_heading(doc, "Dual-Mode Test Flow", 2)
add_code_block(doc, [
    "ROM firmware starts after reset",
    "read 0x40000000",
    "read 0x40000004",
    "read 0x40000008",
    "read 0x4000000c",
    "if final word is 0xaabbccdd, write 0xaabbccdd to 0x20008000",
    "program/erase functions are not exercised in this read-comparison benchmark",
])

add_heading(doc, "Dual-Mode Timing Metrics", 2)
add_kv_table(doc, [
    ("Init time", "Reset release to the first direct QSPI read completion"),
    ("Operation time", "First QSPI read completion to fourth QSPI read completion"),
    ("Total read time", "Reset release to fourth read completion"),
    ("Marker latency", "Fourth read completion to final marker write"),
])

add_heading(doc, "How To Interpret PPA Results", 1)
add_matrix_table(
    doc,
    ["Metric", "Legacy SPI + SRAM Shadowing", "Pure QSPI XIP", "Dual-Mode QSPI"],
    [
        (
            "Performance",
            "Pays upfront copy cost, then executes from SRAM.",
            "Avoids copy cost but every flash read pays QSPI access latency.",
            "Similar read path to pure XIP, plus possible overhead from richer control logic.",
        ),
        (
            "Power",
            "SPI and SRAM active during shadowing; runtime can avoid flash for shadowed code.",
            "QSPI flash path remains active for direct reads.",
            "Read power similar in concept to pure XIP, but more logic can toggle.",
        ),
        (
            "Area",
            "SPI controller plus SRAM dependency.",
            "Read-only QSPI controller, usually simpler than dual mode.",
            "Largest feature set: XIP read, write enable, program, erase, XIP lock/drop logic.",
        ),
    ],
    [1.2, 1.75, 1.75, 1.8],
)

add_heading(doc, "Important Fairness Notes", 1)
add_body(
    doc,
    "The architectures are intentionally not doing identical micro-operations. Legacy performs shadowing because that is its architecture. "
    "Pure XIP and dual-mode QSPI do not shadow into SRAM because their purpose is to replace the shadowing method with direct flash access."
)
add_body(
    doc,
    "The common payload and marker keep correctness and measurement boundaries consistent, while preserving the real architectural distinction: "
    "legacy measures flash-to-SRAM initialization plus SRAM execution, whereas QSPI designs measure direct flash read behavior."
)
add_body(
    doc,
    "For PPA reporting, use the testbench's printed init and operation times, then pair them with synthesis area and power estimates for each controller. "
    "When comparing pure XIP and dual mode, remember that the current test uses only the read path; the dual-mode area/power still includes unused write/erase capability."
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("Flash Controller Architecture Reference").font.size = Pt(9)

doc.save(OUT)
print(OUT)
